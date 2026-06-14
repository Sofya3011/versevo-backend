"""Versevo Backend API v5.2 — на основе рабочей версии 5.0
- Eager init переводчиков и HF-анализа (как в рабочем коде)
- Ленивая загрузка только тяжёлых моделей (pipeline)
- Добавлены недостающие endpoint'ы для Flutter
"""

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import logging
import os
import sys
import base64
import uuid
import re
import json
from datetime import datetime
from collections import Counter
from enum import Enum

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)

def is_valid_text(text: str, min_ratio: float = 0.8) -> bool:
    """Проверяет, что строка — это настоящий текст, а не бинарный мусор"""
    if not text or len(text.strip()) < 10:
        return False
    printable = sum(1 for c in text if c.isprintable() or c in '\n\r\t')
    return (printable / max(len(text), 1)) >= min_ratio

# ========== МОДЕЛИ ==========
class TranslateRequest(BaseModel):
    text: str
    target_language: str = "ru"
    source_language: Optional[str] = None
    style: str = "artistic"

class LoginRequest(BaseModel):
    email: str
    password: str

class RegisterRequest(BaseModel):
    email: str
    username: str
    password: str

class AnalysisRequest(BaseModel):
    document_id: int
    analysis_type: str = "general"

class AIAnalysisRequest(BaseModel):
    document_id: int
    analysis_type: str = "full"
    language: str = "ru"

class ChatRequest(BaseModel):
    document_id: int
    question: str
    language: str = "ru"

# ========== APP ==========
app = FastAPI(title="Versevo Backend API", version="5.2.0")

app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True,
                   allow_methods=["*"], allow_headers=["*"])

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

try:
    app.mount("/uploads", StaticFiles(directory=UPLOAD_FOLDER), name="uploads")
except Exception:
    pass

PORT = int(os.getenv("PORT", 8080))

documents_store = {}
current_doc_id = 1

NLTK_AVAILABLE = False

# ========== ЛОКАЛЬНЫЙ ПЕРЕВОДЧИК (как в рабочем коде) ==========
class LocalTranslator:
    def __init__(self):
        logger.info("🚀 Инициализация улучшенного переводчика")
        self.translation_dict = {
            'en-ru': {
                'hello': 'привет', 'world': 'мир', 'book': 'книга',
                'read': 'читать', 'page': 'страница', 'chapter': 'глава',
                'text': 'текст', 'document': 'документ', 'translate': 'переводить',
                'library': 'библиотека', 'author': 'автор', 'title': 'название',
                'content': 'содержание', 'analysis': 'анализ', 'summary': 'краткое содержание',
                'character': 'персонаж', 'plot': 'сюжет', 'story': 'история',
                'novel': 'роман', 'poem': 'стихотворение', 'literature': 'литература',
                'the': '', 'a': '', 'an': '', 'and': 'и', 'or': 'или',
                'but': 'но', 'in': 'в', 'on': 'на', 'at': 'в', 'to': 'к',
                'for': 'для', 'with': 'с', 'from': 'из', 'of': 'из', 'by': 'от',
                'is': 'является', 'are': 'являются', 'was': 'был', 'were': 'были',
                'have': 'иметь', 'has': 'имеет', 'do': 'делать', 'does': 'делает',
                'can': 'мочь', 'could': 'мог', 'will': 'будет', 'would': 'бы',
                'good': 'хороший', 'bad': 'плохой', 'new': 'новый', 'old': 'старый',
                'big': 'большой', 'small': 'маленький', 'beautiful': 'красивый',
                'interesting': 'интересный', 'important': 'важный',
            },
            'ru-en': {
                'привет': 'hello', 'мир': 'world', 'книга': 'book',
                'читать': 'read', 'глава': 'chapter', 'текст': 'text',
                'документ': 'document', 'автор': 'author', 'название': 'title',
                'содержание': 'content', 'анализ': 'analysis',
                'персонаж': 'character', 'сюжет': 'plot', 'история': 'story',
                'литература': 'literature',
                'и': 'and', 'или': 'or', 'но': 'but', 'в': 'in', 'на': 'on',
                'для': 'for', 'с': 'with',
            }
        }
        self.literary_patterns = {
            'en-ru': [
                (r'\bIt is\b', 'Это'), (r'\bhe said\b', 'сказал он'),
                (r'\bshe said\b', 'сказала она'), (r'\bthey said\b', 'сказали они'),
                (r'\bI think\b', 'Я думаю'), (r'\byou know\b', 'знаете ли'),
                (r'\bof course\b', 'конечно'), (r'\bin fact\b', 'на самом деле'),
                (r'\bat first\b', 'сначала'), (r'\bat last\b', 'наконец'),
            ]
        }

    def _dictionary_translate(self, text: str, lang_key: str) -> str:
        words = re.findall(r'\b\w+\b|[^\w\s]', text)
        translated = []
        d = self.translation_dict[lang_key]
        for word in words:
            if re.match(r'^\w+$', word):
                w = word.lower()
                t = d.get(w, word)
                if t:
                    if word[0].isupper():
                        t = t[0].upper() + t[1:] if len(t) > 0 else t
                    translated.append(t)
                else:
                    translated.append('')
            else:
                translated.append(word)
        result = ' '.join(translated)
        result = re.sub(r'\s+([.,!?;:])', r'\1', result)
        result = re.sub(r'\s+', ' ', result).strip()
        return result

    def _apply_style(self, text: str, style: str) -> str:
        return text

    def translate(self, text: str, source_lang: str, target_lang: str, style: str = "artistic") -> str:
        if source_lang == target_lang:
            return self._apply_style(text, style)
        key = f"{source_lang}-{target_lang}"
        if len(text) > 800:
            text = text[:800] + "..."
        if key in self.literary_patterns:
            for pattern, replacement in self.literary_patterns[key]:
                text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        if key in self.translation_dict:
            result = self._dictionary_translate(text, key)
        else:
            result = text
        result = self._apply_style(result, style)
        return result

# ========== HF ПЕРЕВОДЧИК (с Google Translate API как основной) ==========
class HuggingFaceTranslator:
    def __init__(self):
        logger.info("🌍 Инициализация переводчика")
        self.hf_enabled = False
        self.fallback_translator = LocalTranslator()
        self._init_translation()

    def _init_translation(self):
        try:
            import torch
            device_name = "CUDA" if torch.cuda.is_available() else "CPU"
            self._hf_available = False
            self._google_translator = None
            try:
                from deep_translator import GoogleTranslator
                self._google_translator = GoogleTranslator
                logger.info("✅ Google Translate API доступен")
            except ImportError:
                logger.warning("⚠️ deep_translator не установлен")
            self.hf_enabled = True
            logger.info(f"✅ Переводчик инициализирован (устройство: {device_name})")
        except Exception as e:
            logger.error(f"❌ Ошибка инициализации переводчика: {e}")

    def _do_google_translate(self, text: str, source: str, target: str) -> str:
        if not self._google_translator:
            return ""
        try:
            t = self._google_translator(source=source, target=target)
            result = t.translate(text)
            if result and len(result) > 0:
                return result
        except Exception as e:
            logger.warning(f"⚠️ Google Translate error: {e}")
        if len(text) > 4000:
            chunks = _split_sentences(text, 3000)
            parts = []
            for chunk in chunks:
                try:
                    t = self._google_translator(source=source, target=target)
                    r = t.translate(chunk)
                    if r:
                        parts.append(r)
                except Exception:
                    pass
            if parts:
                return ' '.join(parts)
        return ""

    def translate(self, text: str, source_lang: str, target_lang: str, style: str = "artistic") -> str:
        if source_lang == target_lang:
            return self.fallback_translator._apply_style(text, style)
        supported_pairs = ['en-ru', 'ru-en']
        key = f"{source_lang}-{target_lang}"
        if key not in supported_pairs:
            logger.warning(f"⚠️ Неподдерживаемая пара переводов: {key}")
            return self.fallback_translator.translate(text, source_lang, target_lang, style)
        logger.info(f"🔄 Перевод {len(text)} символов: {source_lang} → {target_lang}")
        try:
            if len(text) <= 3000:
                result = self._do_google_translate(text, source_lang, target_lang)
                if result:
                    logger.info(f"✅ Перевод: {len(text)} → {len(result)} символов")
                    result = self.fallback_translator._apply_style(result, style)
                    return result
            else:
                chunks = _split_sentences(text, 2000)
                translated_parts = []
                for i, chunk in enumerate(chunks):
                    r = self._do_google_translate(chunk, source_lang, target_lang)
                    if r:
                        translated_parts.append(r)
                    else:
                        translated_parts.append(self.fallback_translator.translate(chunk, source_lang, target_lang, style))
                if translated_parts:
                    result = ' '.join(translated_parts)
                    result = self.fallback_translator._apply_style(result, style)
                    return result
        except Exception as e:
            logger.warning(f"⚠️ Google Translate error: {e}")
        logger.warning(f"⚠️ Использую fallback перевод для {key}")
        return self.fallback_translator.translate(text, source_lang, target_lang, style)

    def is_available(self, source_lang: str, target_lang: str) -> bool:
        return self.hf_enabled

# ========== ИНИЦИАЛИЗАЦИЯ (как в рабочем коде — eager) ==========
logger.info("🚀 ИНИЦИАЛИЗАЦИЯ VERSION 5.2")
local_translator = LocalTranslator()
hf_translator = HuggingFaceTranslator()

# ========== HF АНАЛИЗ (как в рабочем коде) ==========
HUGGING_FACE_ENABLED = False
HF_ANALYSIS_PIPELINES = {}

def init_huggingface_for_analysis():
    global HUGGING_FACE_ENABLED, HF_ANALYSIS_PIPELINES
    try:
        import transformers
        import torch
        device = 0 if torch.cuda.is_available() else -1
        device_name = "CUDA" if torch.cuda.is_available() else "CPU"
        logger.info(f"🏗️ Инициализация Hugging Face анализа на {device_name}")
        models = {
            "sentiment": {"model_name": "blanchefort/rubert-base-cased-sentiment", "task": "sentiment-analysis"},
            "summarization": {"model_name": "IlyaGusev/rut5_base_sum_gazeta", "task": "summarization"},
            "ner": {"model_name": "Babelscape/wikineural-multilingual-ner", "task": "ner"},
        }
        HF_ANALYSIS_PIPELINES = {
            task: {"model_name": cfg["model_name"], "pipeline": None, "task": cfg["task"]}
            for task, cfg in models.items()
        }
        HUGGING_FACE_ENABLED = True
        logger.info("✅ Hugging Face анализ инициализирован (ленивая загрузка)")
        try:
            logger.info("🔄 Предзагрузка модели анализа тональности...")
            from transformers import pipeline
            HF_ANALYSIS_PIPELINES["sentiment"]["pipeline"] = pipeline(
                "sentiment-analysis", model=models["sentiment"]["model_name"], device=device
            )
            logger.info("✅ Модель анализа тональности загружена")
        except Exception as e:
            logger.warning(f"⚠️ Не удалось загрузить sentiment модель: {e}")
    except ImportError as e:
        logger.warning(f"⚠️ transformers не установлен: {e}")
    except Exception as e:
        logger.error(f"❌ Ошибка инициализации Hugging Face анализа: {e}")

init_huggingface_for_analysis()

# ========== NLTK ==========
try:
    import nltk
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords', quiet=True)
    NLTK_AVAILABLE = True
    logger.info("✅ NLTK успешно инициализирован")
except ImportError:
    NLTK_AVAILABLE = False
    logger.warning("⚠️ NLTK не установлен")

def get_hf_analysis_pipeline(task: str):
    if task not in HF_ANALYSIS_PIPELINES:
        return None
    if HF_ANALYSIS_PIPELINES[task]["pipeline"] is None and HUGGING_FACE_ENABLED:
        try:
            from transformers import pipeline
            import torch
            device = 0 if torch.cuda.is_available() else -1
            cfg = HF_ANALYSIS_PIPELINES[task]
            logger.info(f"🔄 Загрузка модели анализа {cfg['model_name']}...")
            if task == "summarization":
                pipe = pipeline("summarization", model=cfg["model_name"],
                              tokenizer=cfg["model_name"], device=device,
                              max_length=150, min_length=50)
            elif task == "ner":
                pipe = pipeline("ner", model=cfg["model_name"], device=device,
                              grouped_entities=True, ignore_labels=["O"])
            else:
                pipe = pipeline(cfg["task"], model=cfg["model_name"], device=device)
            HF_ANALYSIS_PIPELINES[task]["pipeline"] = pipe
            logger.info(f"✅ Модель анализа {cfg['model_name']} загружена")
        except Exception as e:
            logger.error(f"❌ Ошибка загрузки модели {task}: {e}")
            return None
    return HF_ANALYSIS_PIPELINES[task]["pipeline"]

# ========== УТИЛИТЫ ==========
def extract_text_from_file(file_path: str, file_type: str) -> str:
    try:
        if file_type == 'pdf':
            try:
                import fitz
                doc = fitz.open(file_path)
                pages = []
                for page in doc:
                    t = page.get_text().strip()
                    if t:
                        pages.append(t)
                doc.close()
                if pages:
                    return "\n\n".join(pages)
            except Exception:
                pass
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract
                t = pdfminer_extract(file_path)
                if t and t.strip():
                    return t.strip()
            except Exception:
                pass
            return ""
        elif file_type in ['docx', 'doc']:
            try:
                import docx
                doc = docx.Document(file_path)
                paras = [p.text for p in doc.paragraphs if p.text.strip()]
                if paras:
                    return "\n\n".join(paras)
                tables = []
                for table in doc.tables:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells:
                            tables.append(" | ".join(cells))
                if tables:
                    return "\n\n".join(tables)
                return ""
            except Exception:
                return ""
        elif file_type == 'txt':
            try:
                with open(file_path, "r", encoding='utf-8', errors='ignore') as f:
                    return f.read()
            except Exception:
                pass
            try:
                import chardet
                with open(file_path, "rb") as f:
                    raw = f.read()
                    detected = chardet.detect(raw)
                    enc = detected.get("encoding", "utf-8") or "utf-8"
                    return raw.decode(enc, errors="replace")
            except Exception:
                pass
            try:
                with open(file_path, "r", encoding='cp1251', errors='replace') as f:
                    return f.read()
            except Exception:
                return ""
        elif file_type == 'epub':
            try:
                import zipfile
                with zipfile.ZipFile(file_path, 'r') as z:
                    html_files = [f for f in z.namelist() if f.endswith(('.xhtml', '.html', '.htm'))]
                    if not html_files:
                        html_files = [f for f in z.namelist() if '.' not in f.split('/')[-1] or f.endswith(('.xhtml', '.html', '.htm'))]
                    html_files.sort()
                    from html.parser import HTMLParser
                    class TextExtractor(HTMLParser):
                        def __init__(self):
                            super().__init__()
                            self.text_parts = []
                        def handle_data(self, data):
                            stripped = data.strip()
                            if stripped:
                                self.text_parts.append(stripped)
                    texts = []
                    for hf in html_files:
                        try:
                            content = z.read(hf)
                            try:
                                decoded = content.decode('utf-8')
                            except UnicodeDecodeError:
                                decoded = content.decode('utf-8', errors='replace')
                            parser = TextExtractor()
                            parser.feed(decoded)
                            chunk = ' '.join(parser.text_parts)
                            if chunk.strip():
                                texts.append(chunk.strip())
                        except Exception:
                            pass
                    if texts:
                        return "\n\n".join(texts)
                    return ""
            except Exception:
                pass
            try:
                import ebooklib
                from ebooklib import epub
                book = epub.read_epub(file_path)
                texts = []
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        content = item.get_content()
                        try:
                            decoded = content.decode('utf-8')
                        except UnicodeDecodeError:
                            decoded = content.decode('utf-8', errors='replace')
                        from html.parser import HTMLParser
                        class TE(HTMLParser):
                            def __init__(self):
                                super().__init__()
                                self.parts = []
                            def handle_data(self, d):
                                s = d.strip()
                                if s:
                                    self.parts.append(s)
                        parser = TE()
                        parser.feed(decoded)
                        chunk = ' '.join(parser.parts)
                        if chunk.strip():
                            texts.append(chunk.strip())
                if texts:
                    return "\n\n".join(texts)
                return ""
            except Exception:
                return ""
        return ""
    except Exception:
        return ""

def detect_language_safe(text: str) -> str:
    if not text or len(text.strip()) < 10:
        return "en"
    cyrillic = sum(1 for c in text if 'а' <= c <= 'я' or 'А' <= c <= 'Я')
    latin = sum(1 for c in text if 'a' <= c <= 'z' or 'A' <= c <= 'Z')
    return "ru" if cyrillic > latin * 1.5 else "en"

def _split_sentences(text: str, max_chars: int = 1500) -> List[str]:
    """Разбивает текст на части по границам предложений, не превышая max_chars"""
    if len(text) <= max_chars:
        return [text]
    result = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        if end < len(text):
            cut = text.rfind('. ', start, end)
            if cut < start:
                cut = text.rfind('! ', start, end)
            if cut < start:
                cut = text.rfind('? ', start, end)
            if cut < start:
                cut = text.rfind('\n\n', start, end)
            if cut < start:
                cut = end
            else:
                cut += 1
        else:
            cut = end
        chunk = text[start:cut].strip()
        if chunk:
            result.append(chunk)
        start = cut
    if not result:
        result = [text[:max_chars]]
    return result

def detect_chapters(text: str) -> List[Dict]:
    chapters = []
    if not text:
        return [{'title': 'Документ', 'content': 'Нет содержимого'}]
    text = re.sub(r'\n{3,}', '\n\n', text.strip())
    patterns = [
        r'^\s*(?:(?:ГЛАВА|Глава|Гл\.?|CHAPTER|Chapter|Ch\.?|SECTION|Section|Sec\.?|РАЗДЕЛ|Раздел|PART|Part|ЧАСТЬ|Часть)\s*[\.:\s]*\s*[IVXLCDM\d]+[\.,:\s]*.*)$',
        r'^\s*(?:[IVXLCDM]+[\.\)]\s+.*)$',
        r'^\s*\d+[\.\)]\s+[A-ZА-Я].*$',
        r'^\s*\d{1,2}\.\d{1,2}\s+.*$',
        r'^\s*[A-ZА-Я][A-ZА-Я\s]{2,}[\.\?!]?$',
        r'^\s*.+\n[-=]{3,}$',
        r'^\s*PROLOGUE|Prologue|EPILOGUE|Epilogue|FOREWORD|Foreword|INTRODUCTION|Introduction|ПРОЛОГ|Пролог|ЭПИЛОГ|Эпилог|ВВЕДЕНИЕ|Введение|ПРЕДИСЛОВИЕ|Предисловие$',
    ]
    paragraphs = text.split('\n\n')
    current_chapter = None
    chapter_content = []
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        is_title = any(re.match(p, paragraph, re.MULTILINE | re.IGNORECASE) for p in patterns)
        if is_title:
            if current_chapter is not None and chapter_content:
                chapters.append({'title': current_chapter, 'content': '\n\n'.join(chapter_content)})
            current_chapter = paragraph[:120]
            chapter_content = []
        else:
            if current_chapter is None:
                current_chapter = 'Начало'
            chapter_content.append(paragraph)
    if current_chapter and chapter_content:
        chapters.append({'title': current_chapter, 'content': '\n\n'.join(chapter_content)})
    if not chapters:
        for i in range(0, len(text), 5000):
            chunk = text[i:i + 5000]
            if chunk.strip():
                chapters.append({'title': f'Часть {len(chapters) + 1}', 'content': chunk})
    return chapters

# ========== HEALTH ==========
@app.get("/")
async def root():
    return {
        "message": "Versevo Backend API v5.2",
        "version": "5.2.0",
        "status": "running",
        "translation": {
            "huggingface_available": hf_translator.is_available('en', 'ru'),
            "fallback_available": True,
        },
        "analysis": {
            "huggingface_available": HUGGING_FACE_ENABLED,
            "nltk_available": NLTK_AVAILABLE,
        },
        "timestamp": datetime.now().isoformat(),
    }

@app.get("/api/health")
@app.get("/api/flutter/health")
async def health_check():
    return {
        "status": "healthy",
        "service": "versevo-backend",
        "version": "5.2.0",
        "translation": {
            "huggingface": hf_translator.is_available('en', 'ru'),
            "fallback": True,
        },
        "analysis": {
            "huggingface": HUGGING_FACE_ENABLED,
            "nltk": NLTK_AVAILABLE,
        },
        "timestamp": datetime.now().isoformat(),
    }

@app.get("/api/warmup")
async def warmup():
    return {"status": "warm", "timestamp": datetime.now().isoformat()}

@app.get("/api/analyze/ai/health")
async def ai_health_check():
    return {
        "status": "healthy" if HUGGING_FACE_ENABLED else "unavailable",
        "service": "huggingface",
        "available": HUGGING_FACE_ENABLED,
        "models_loaded": [k for k, v in HF_ANALYSIS_PIPELINES.items() if v["pipeline"] is not None],
        "models_available": list(HF_ANALYSIS_PIPELINES.keys()),
        "timestamp": datetime.now().isoformat(),
    }

# ========== ДОКУМЕНТЫ ==========
@app.post("/api/documents/upload-base64")
async def upload_document_base64(request: dict):
    global current_doc_id
    try:
        filename = request.get("filename", "unknown.txt")
        file_data = request.get("file_data", "")
        if not file_data:
            raise HTTPException(status_code=400, detail="No file data provided")
        content_bytes = base64.b64decode(file_data)
        file_id = str(uuid.uuid4())
        ext = filename.split('.')[-1].lower() if '.' in filename else 'txt'
        file_path = f"{UPLOAD_FOLDER}/{file_id}.{ext}"
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        with open(file_path, "wb") as f:
            f.write(content_bytes)
        content_str = extract_text_from_file(file_path, ext)
        if content_str and not is_valid_text(content_str):
            logger.warning(f"⚠️ Извлечённый текст из {filename} похож на бинарные данные, отбрасываю")
            content_str = ""
        if not content_str or content_str.strip() == "":
            logger.warning(f"⚠️ Не удалось извлечь текст из {filename}")
            content_str = f"Документ: {filename}\nТип: {ext}\n\nСодержимое недоступно для автоматического извлечения."
        language = detect_language_safe(content_str)
        chapters = detect_chapters(content_str)
        word_count = len(content_str.split())
        char_count = len(content_str)
        reading_time = max(1, word_count // 200)
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        doc = {
            "id": current_doc_id, "filename": filename,
            "original_filename": filename, "content": content_str,
            "language": language, "file_type": ext, "file_path": file_path,
            "file_id": file_id, "word_count": word_count,
            "char_count": char_count, "chapter_count": len(chapters),
            "reading_time_minutes": reading_time, "file_size": file_size,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(), "chapters": chapters,
        }
        documents_store[current_doc_id] = doc
        current_doc_id += 1
        logger.info(f"✅ Документ загружен: {filename} (ID: {doc['id']}, слов: {word_count})")
        return {
            "id": doc["id"], "filename": doc["filename"],
            "language": doc["language"], "file_type": doc["file_type"],
            "word_count": doc["word_count"], "char_count": doc["char_count"],
            "chapter_count": doc["chapter_count"],
            "reading_time_minutes": doc["reading_time_minutes"],
            "file_size": doc["file_size"],
            "created_at": doc["created_at"],
            "content_preview": doc["content"][:300] + "..." if len(doc["content"]) > 300 else doc["content"],
        }
    except Exception as e:
        logger.error(f"❌ Upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.get("/api/documents")
async def get_documents():
    docs = list(documents_store.values())
    return [{
        "id": d["id"], "filename": d["filename"],
        "language": d["language"], "file_type": d["file_type"],
        "word_count": d["word_count"], "char_count": d["char_count"],
        "chapter_count": d["chapter_count"],
        "reading_time_minutes": d["reading_time_minutes"],
        "file_size": d.get("file_size", 0),
        "created_at": d["created_at"],
        "updated_at": d.get("updated_at", d["created_at"]),
    } for d in sorted(docs, key=lambda x: x["created_at"], reverse=True)]

@app.get("/api/documents/{document_id}")
async def get_document(document_id: int):
    if document_id not in documents_store:
        raise HTTPException(status_code=404, detail="Document not found")
    d = documents_store[document_id]
    return {
        "id": d["id"], "filename": d["filename"], "content": d["content"],
        "language": d["language"], "file_type": d["file_type"],
        "word_count": d["word_count"], "char_count": d["char_count"],
        "chapter_count": d["chapter_count"],
        "reading_time_minutes": d["reading_time_minutes"],
        "file_size": d.get("file_size", 0),
        "created_at": d["created_at"],
        "updated_at": d.get("updated_at", d["created_at"]),
        "chapters": d["chapters"],
    }

@app.delete("/api/documents/{document_id}")
async def delete_document(document_id: int):
    if document_id not in documents_store:
        raise HTTPException(status_code=404, detail="Document not found")
    doc = documents_store[document_id]
    if os.path.exists(doc["file_path"]):
        try:
            os.remove(doc["file_path"])
        except:
            pass
    del documents_store[document_id]
    return {"status": "success", "deleted_id": document_id}

# ========== ПЕРЕВОД ==========
@app.post("/api/translate/text")
async def translate_text(request: TranslateRequest):
    try:
        if not request.text or len(request.text.strip()) == 0:
            raise HTTPException(status_code=400, detail="Text is empty")
        source_lang = request.source_language
        if not source_lang or source_lang == "auto":
            source_lang = detect_language_safe(request.text)
        target_lang = request.target_language
        use_huggingface = hf_translator.is_available(source_lang, target_lang)
        if use_huggingface:
            logger.info(f"🔄 Используем Hugging Face перевод: {source_lang} → {target_lang}")
            translated_text = hf_translator.translate(request.text, source_lang, target_lang, request.style)
            translation_service = "huggingface"
        else:
            logger.info(f"🔄 Используем fallback перевод: {source_lang} → {target_lang}")
            translated_text = local_translator.translate(request.text, source_lang, target_lang, request.style)
            translation_service = "fallback"
        return {
            "original_text": request.text,
            "translated_text": translated_text,
            "source_language": source_lang,
            "target_language": target_lang,
            "style": request.style,
            "translation_service": translation_service,
            "original_length": len(request.text),
            "translated_length": len(translated_text),
            "huggingface_used": use_huggingface,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Translate error: {e}")
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.post("/api/translate/document/{document_id}")
async def translate_document(document_id: int, request: dict):
    global current_doc_id
    try:
        if document_id not in documents_store:
            raise HTTPException(status_code=404, detail="Document not found")
        doc = documents_store[document_id]
        target_lang = request.get("target_language", "ru")
        source_lang = request.get("source_language") or doc.get("language", "auto")
        logger.info(f"🔄 Перевод документа {document_id} ({doc['filename']}) на {target_lang}")

        content = doc.get("content", "")
        if not content or len(content.strip()) < 10:
            raise HTTPException(status_code=400, detail="Документ пуст или не содержит текста для перевода")

        paragraphs = content.split('\n\n')
        translated_paragraphs = []
        hf = hf_translator
        local = local_translator
        total = len(paragraphs)

        for i, para in enumerate(paragraphs):
            para = para.strip()
            if not para:
                translated_paragraphs.append('')
                continue
            if len(para) <= 2000:
                logger.info(f"🔄 Перевод абзаца {i + 1}/{total} ({len(para)} символов)")
                use_hf = hf.is_available(source_lang if source_lang != "auto" else "en", target_lang)
                if use_hf:
                    translated = hf.translate(para, source_lang if source_lang != "auto" else detect_language_safe(para), target_lang)
                else:
                    translated = local.translate(para, source_lang if source_lang != "auto" else detect_language_safe(para), target_lang)
                translated_paragraphs.append(translated)
            else:
                chunks = _split_sentences(para, 1500)
                para_translated = []
                for j, chunk in enumerate(chunks):
                    logger.info(f"🔄 Перевод абзаца {i + 1}/{total}, часть {j + 1}/{len(chunks)} ({len(chunk)} символов)")
                    use_hf = hf.is_available(source_lang if source_lang != "auto" else "en", target_lang)
                    if use_hf:
                        translated = hf.translate(chunk, source_lang if source_lang != "auto" else detect_language_safe(chunk), target_lang)
                    else:
                        translated = local.translate(chunk, source_lang if source_lang != "auto" else detect_language_safe(chunk), target_lang)
                    para_translated.append(translated)
                translated_paragraphs.append(' '.join(para_translated))

        translated_content = '\n\n'.join(translated_paragraphs)
        new_id = current_doc_id + 1
        import copy
        td = copy.deepcopy(doc)
        td["id"] = new_id
        td["filename"] = f"[{target_lang.upper()}] {doc['filename']}"
        td["language"] = target_lang
        td["content"] = translated_content
        td["created_at"] = datetime.now().isoformat()
        td["updated_at"] = datetime.now().isoformat()
        chapters = detect_chapters(translated_content)
        td["chapters"] = chapters
        td["chapter_count"] = len(chapters)
        td["word_count"] = len(translated_content.split())
        td["char_count"] = len(translated_content)
        documents_store[new_id] = td
        current_doc_id = new_id + 1
        logger.info(f"✅ Документ {document_id} переведён на {target_lang} -> ID {new_id} ({len(translated_content)} символов)")
        return {"status": "success", "translated_document_id": new_id, "message": "Документ переведен"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Translate document error: {e}")
        raise HTTPException(status_code=500, detail=f"Document translation failed: {str(e)}")

# ========== БАЗОВЫЙ АНАЛИЗ ==========
def _perform_basic_analysis(text: str) -> Dict[str, Any]:
    result = {"summary": "", "themes": [], "sentiment": "Нейтральный",
              "complexity": "Средний", "key_points": [], "statistics": {},
              "language_features": {}}
    if not text or len(text.strip()) < 10:
        result["summary"] = "Текст слишком короткий для анализа"
        return result
    try:
        words = [w for w in text.split() if w.strip()]
        sentences = re.split(r'[.!?]+', text)
        paragraphs = text.split('\n\n')
        sentences = [s.strip() for s in sentences if s.strip()]
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        wc = len(words)
        sc = len(sentences)
        pc = len(paragraphs)
        avg = wc / sc if sc > 0 else 0
        complexity = "Простой" if avg < 8 else "Средний" if avg < 15 else "Сложный"
        result["complexity"] = complexity
        if sc >= 3:
            summary_list = [s.strip() for s in sentences[:4] if len(s.strip()) > 10 and not s.strip().isupper()]
            if summary_list:
                result["summary"] = " ".join(summary_list)
                if len(result["summary"]) > 250:
                    result["summary"] = result["summary"][:250] + "..."
            else:
                result["summary"] = text[:200] + "..." if len(text) > 200 else text
        else:
            result["summary"] = text[:200] + "..." if len(text) > 200 else text
        proper = re.findall(r'\b[A-Z][a-z]+\b', text[:1000])
        themes = []
        if proper:
            c = Counter(n.lower() for n in proper)
            themes = [n for n, _ in c.most_common(5) if c[n] > 1 and len(n) > 3][:3]
        if not themes:
            word_freq = Counter(w.lower() for w in words if len(w) > 3)
            stop = {'это','что','как','для','того','чтобы','если','когда','или','и','но','а',
                    'the','and','but','for','with','from','that','this','was','were'}
            themes = [w for w, _ in word_freq.most_common(10) if w not in stop and word_freq[w] > 1][:3]
        result["themes"] = themes[:3] if themes else ["Документ", "Текст", "Содержание"]
        result["sentiment"] = "Нейтральный"
        result["statistics"] = {
            "word_count": wc, "sentence_count": sc, "paragraph_count": pc,
            "avg_sentence_length": round(avg, 1),
            "avg_word_length": round(sum(len(w) for w in words) / wc if wc > 0 else 0, 1),
            "reading_time_minutes": max(1, wc // 200),
        }
        cyrillic = sum(1 for c in text if 'а' <= c <= 'я' or 'А' <= c <= 'Я')
        latin = sum(1 for c in text if 'a' <= c <= 'z' or 'A' <= c <= 'Z')
        result["language_features"] = {
            "detected_language": "ru" if cyrillic > latin else "en",
            "has_dialogue": bool(re.search(r'["\'«»]', text)),
            "has_numbers": bool(re.search(r'\d+', text)),
            "has_questions": "?" in text,
            "has_exclamations": "!" in text,
        }
        result["key_points"] = [
            f"Объем: {wc} слов", f"Сложность: {complexity}",
            f"Время чтения: {max(1, wc // 200)} мин",
        ]
        if themes:
            result["key_points"].append(f"Темы: {', '.join(themes[:2])}")
    except Exception as e:
        logger.error(f"Ошибка базового анализа: {e}")
        result["summary"] = "Произошел сбой при анализе текста"
        result["key_points"] = ["Не удалось выполнить анализ"]
    return result

@app.post("/api/analyze")
async def analyze_document(request: AnalysisRequest):
    try:
        if request.document_id not in documents_store:
            raise HTTPException(status_code=404, detail="Document not found")
        doc = documents_store[request.document_id]
        content = doc["content"]
        if not content or len(content.strip()) < 10:
            raise HTTPException(status_code=400, detail="Document has no content")
        logger.info(f"🔍 Базовый анализ документа {request.document_id}")
        r = _perform_basic_analysis(content)
        return {
            "document_id": request.document_id, "filename": doc["filename"],
            "language": doc["language"],
            "summary": r["summary"], "themes": r["themes"],
            "sentiment": r["sentiment"], "complexity": r["complexity"],
            "writing_style": "Информационный",
            "key_points": r["key_points"],
            "statistics": r["statistics"],
            "language_features": r["language_features"],
            "document_statistics": {
                "word_count": doc["word_count"], "char_count": doc["char_count"],
                "chapter_count": doc["chapter_count"],
                "reading_time_minutes": doc["reading_time_minutes"],
                "file_type": doc["file_type"],
            },
            "characters": [], "entities": [],
            "ai_analysis": False, "fallback": False,
            "analysis_type": request.analysis_type,
            "analysis_timestamp": datetime.now().isoformat(),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Ошибка базового анализа: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

# ========== AI АНАЛИЗ (как в рабочем коде) ==========
def _perform_ai_analysis(text: str) -> Dict[str, Any]:
    result = {"summary": "", "themes": [], "sentiment": "Нейтральный",
              "writing_style": "Информационный", "key_points": [],
              "entities": [], "ai_analysis": False, "fallback": False, "models_used": []}
    if not HUGGING_FACE_ENABLED or not text or len(text.strip()) < 50:
        result["fallback"] = True
        return result
    if not is_valid_text(text):
        logger.warning("⚠️ AI анализ пропущен: текст содержит бинарные данные")
        result["fallback"] = True
        return result
    try:
        sample_text = text[:2000]
        sentiment_pipe = get_hf_analysis_pipeline("sentiment")
        if sentiment_pipe:
            try:
                sr = sentiment_pipe(sample_text[:512])
                if sr and len(sr) > 0:
                    label = sr[0].get("label", "NEUTRAL").upper()
                    sm = {"POSITIVE": "Положительный", "NEGATIVE": "Отрицательный", "NEUTRAL": "Нейтральный",
                          "LABEL_0": "Отрицательный", "LABEL_1": "Нейтральный", "LABEL_2": "Положительный"}
                    result["sentiment"] = sm.get(label, "Нейтральный")
                    result["models_used"].append("sentiment")
                    result["ai_analysis"] = True
            except Exception as e:
                logger.warning(f"⚠️ Sentiment error: {e}")
        if len(text.split()) > 150:
            summ_pipe = get_hf_analysis_pipeline("summarization")
            if summ_pipe:
                try:
                    clean_text = re.sub(r'\s+', ' ', sample_text.strip())
                    if len(clean_text) > 100:
                        sr2 = summ_pipe(clean_text, max_length=120, min_length=60, do_sample=False)
                        if sr2:
                            summary = sr2[0].get("summary_text", "")
                            summary = re.sub(r'^\[ПЕРЕВОД\]\s*', '', summary)
                            result["summary"] = summary
                            result["models_used"].append("summarization")
                            result["ai_analysis"] = True
                except Exception as e:
                    logger.warning(f"⚠️ Summarization error: {e}")
        ner_pipe = get_hf_analysis_pipeline("ner")
        if ner_pipe:
            try:
                ner_result = ner_pipe(sample_text[:1000])
                entities = []
                for e in ner_result:
                    if isinstance(e, dict):
                        word = e.get("word", "")
                        group = e.get("entity_group", "")
                        if group in ["PER", "ORG", "LOC"] and len(word) > 2 and not re.match(r'^\d+$', word):
                            entities.append({"entity": group, "word": word, "score": round(e.get("score", 0), 3)})
                if entities:
                    unique_entities = []
                    seen = set()
                    for e in entities:
                        key = f"{e['entity']}_{e['word'].lower()}"
                        if key not in seen:
                            unique_entities.append(e)
                            seen.add(key)
                    result["entities"] = unique_entities[:10]
                    result["models_used"].append("ner")
                    result["ai_analysis"] = True
                    et = Counter(e["entity"] for e in result["entities"])
                    themes = []
                    for et_type, _ in et.most_common(3):
                        if et_type == "PER": themes.append("Персонажи")
                        elif et_type == "ORG": themes.append("Организации")
                        elif et_type == "LOC": themes.append("Места")
                    if themes:
                        result["themes"] = themes
            except Exception as e:
                logger.warning(f"⚠️ NER error: {e}")
        wc = len(text.split())
        sc = len(re.split(r'[.!?]+', text))
        ws = "Академический" if wc > 5000 else "Литературный" if sc > 0 and wc / sc > 25 else "Информационный"
        if "?" in text and "!" in text and '"' in text:
            ws = "Диалогический"
        result["writing_style"] = ws
        kp = []
        if result["entities"]:
            people = [e["word"] for e in result["entities"] if e["entity"] == "PER"]
            if people:
                kp.append(f"Персонажи: {', '.join(people[:2])}")
        if result["sentiment"] != "Нейтральный":
            kp.append(f"Тональность: {result['sentiment']}")
        kp.append(f"Стиль письма: {ws}")
        kp.append(f"Время чтения: {max(1, wc // 200)} мин")
        kp.append(f"Объем: {wc} слов")
        result["key_points"] = kp if kp else [f"Объем: {wc} слов"]
        if not result["summary"]:
            sentences = re.split(r'[.!?]+', text)
            summary_list = [s.strip() for s in sentences[:3] if len(s.strip()) > 10]
            if summary_list:
                result["summary"] = " ".join(summary_list)[:300]
            else:
                result["summary"] = text[:200] + "..." if len(text) > 200 else text
        if not result["themes"]:
            nouns = re.findall(r'\b[A-Z][a-z]+\b', text[:1000])
            if nouns:
                cn = Counter(n.lower() for n in nouns)
                common = [n for n, _ in cn.most_common(5) if cn[n] > 1 and len(n) > 3][:3]
                if common:
                    result["themes"] = common
        if not result["themes"]:
            result["themes"] = ["Литература", "Текст", "Содержание"]
        return result
    except Exception as e:
        logger.error(f"❌ Ошибка AI анализа: {e}")
        result["fallback"] = True
        return result

@app.post("/api/analyze/ai/document")
async def analyze_with_ai(request: AIAnalysisRequest):
    try:
        did = request.document_id
        if did not in documents_store:
            raise HTTPException(status_code=404, detail="Document not found")
        doc = documents_store[did]
        content = doc["content"]
        if not content or len(content.strip()) < 10:
            raise HTTPException(status_code=400, detail="Document has no content")
        logger.info(f"🔍 AI анализ документа {did}")
        basic = _perform_basic_analysis(content)
        ai = _perform_ai_analysis(content)
        result = {
            "document_id": did, "filename": doc["filename"], "language": doc["language"],
            "summary": ai.get("summary") or basic["summary"],
            "themes": ai.get("themes") or basic["themes"],
            "sentiment": ai.get("sentiment") or basic["sentiment"],
            "writing_style": ai.get("writing_style") or "Информационный",
            "key_points": ai.get("key_points") or basic["key_points"],
            "entities": ai.get("entities", []),
            "statistics": basic["statistics"],
            "language_features": basic["language_features"],
            "ai_analysis": ai.get("ai_analysis", False),
            "fallback": ai.get("fallback", True),
            "models_used": ai.get("models_used", []),
            "analysis_type": request.analysis_type,
            "analysis_timestamp": datetime.now().isoformat(),
        }
        if ai.get("fallback"):
            result["analysis_notes"] = ["Использован базовый анализ из-за недоступности AI"]
        logger.info(f"✅ AI анализ завершен для документа {did}")
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ AI анализ error: {e}")
        return {
            "document_id": request.document_id,
            "summary": "Произошла ошибка при AI-анализе.",
            "themes": ["Ошибка анализа"], "sentiment": "Не определена",
            "writing_style": "Не определен",
            "key_points": ["Не удалось выполнить полный AI-анализ"],
            "entities": [], "ai_analysis": False, "fallback": True,
            "analysis_timestamp": datetime.now().isoformat(),
        }

@app.post("/api/analyze/gemini/document")
async def analyze_with_gemini(request: AIAnalysisRequest):
    try:
        did = request.document_id
        if did not in documents_store:
            raise HTTPException(status_code=404, detail="Document not found")
        doc = documents_store[did]
        content = doc.get("content", "")
        basic = _perform_basic_analysis(content)
        return {
            "document_id": did, "filename": doc["filename"], "language": doc["language"],
            "summary": basic["summary"], "themes": basic["themes"],
            "sentiment": basic["sentiment"], "writing_style": "Информационный",
            "key_points": basic["key_points"], "entities": [],
            "statistics": basic["statistics"],
            "language_features": basic["language_features"],
            "ai_analysis": False, "fallback": True, "models_used": [],
            "analysis_timestamp": datetime.now().isoformat(),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Gemini analysis error: {e}")
        return {
            "document_id": request.document_id,
            "summary": "Ошибка Gemini-анализа.", "themes": [],
            "sentiment": "Не определена", "writing_style": "Не определен",
            "key_points": ["Ошибка Gemini-анализа"], "entities": [],
            "ai_analysis": False, "fallback": True,
            "analysis_timestamp": datetime.now().isoformat(),
        }

# ========== ЧАТ ==========
def _generate_chat_answer(document: dict, question: str) -> str:
    content = document.get("content", "")
    title = document.get("filename", "документ")
    q = question.lower()
    if not content or len(content.strip()) < 10:
        return f"Документ «{title}» пуст или недоступен для чтения."
    words = content.split()
    wc = len(words)
    sentences = re.split(r'(?<=[.!?])\s+', content)
    if any(w in q for w in ['о чём', 'о чем', 'суть', 'содержание', 'кратко']):
        key = [s.strip() for s in sentences if 30 < len(s.strip()) < 300][:4]
        if key:
            return f"📄 Документ «{title}» — {wc} слов.\n\nКраткое содержание:\n\n" + "\n\n".join(f"{i+1}. {s}" for i, s in enumerate(key))
        return f"📄 Документ «{title}» содержит {wc} слов.\n\n{content[:400]}..."
    if any(w in q for w in ['тезис', 'главн', 'основн', 'темы', 'иде']):
        key = [s.strip() for s in sentences if 40 < len(s.strip()) < 250][:6]
        if len(key) >= 2:
            return f"📌 Основные тезисы «{title}»:\n\n" + "\n\n".join(f"{i+1}. {s}" for i, s in enumerate(key))
    if any(w in q for w in ['персонаж', 'герой', 'кто', 'люди', 'человек']):
        names = re.findall(r'\b[А-Я][а-я]+\b', content[:3000])
        unique = list(dict.fromkeys(n for n in names if len(n) > 1 and n not in ['Это','Что','Как','Для','Они','Мы','Вы','Она','Он','Глава']))
        if unique:
            return f"👥 В документе «{title}» упоминаются: {', '.join(unique[:10])}."
    if any(w in q for w in ['найди', 'найти', 'поиск', 'абзац', 'где']):
        search = re.sub(r'(найди|найти|поиск|абзац|про|где|мне|пожалуйста)\s*', '', question, flags=re.IGNORECASE).strip()
        if len(search) > 2:
            idx = content.lower().find(search.lower())
            if idx != -1:
                start = max(0, idx - 150)
                end = min(len(content), idx + len(search) + 250)
                return f"🔍 По запросу «{search}»:\n\n...{content[start:end]}..."
            return f"🔍 По запросу «{search}» ничего не найдено."
    if any(w in q for w in ['жанр', 'стиль', 'как написан']):
        avg_len = wc / len(sentences) if sentences else 0
        style = 'Литературный / Академический' if avg_len > 25 else 'Художественный / Описательный' if avg_len > 15 else 'Разговорный / Информационный'
        return f"📝 Стиль «{title}»: {style}\nСредняя длина предложения: {avg_len:.1f} слов\nВсего предложений: {len(sentences)}\nОбъём: {wc} слов"
    return (f"📖 Документ «{title}» — {wc} слов.\n\n"
            f"Чтобы получить точный ответ, попробуйте:\n"
            f"• «О чём этот документ?»\n"
            f"• «Выдели основные тезисы»\n"
            f"• «Какие персонажи упоминаются?»\n"
            f"• «Найди абзац про...»")

@app.post("/api/chat/ask")
async def chat_ask(request: ChatRequest):
    try:
        did = request.document_id
        if did not in documents_store:
            raise HTTPException(status_code=404, detail="Document not found")
        doc = documents_store[did]
        answer = _generate_chat_answer(doc, request.question)
        if HUGGING_FACE_ENABLED and len(doc.get("content", "")) > 100 and is_valid_text(doc.get("content", "")):
            try:
                pipe = get_hf_analysis_pipeline("sentiment")
                if pipe and any(w in request.question.lower() for w in ['тональность', 'настроение', 'эмоциональн']):
                    sr = pipe(doc["content"][:512])
                    if sr:
                        label = sr[0].get("label", "NEUTRAL")
                        score = sr[0].get("score", 0.5)
                        sm = {"POSITIVE": "позитивная", "LABEL_2": "позитивная",
                              "NEGATIVE": "негативная", "LABEL_0": "негативная",
                              "NEUTRAL": "нейтральная", "LABEL_1": "нейтральная"}
                        answer += f"\n\nАнализ тональности: {sm.get(label, 'нейтральная')} (уверенность: {score:.0%})"
            except:
                pass
        return {"answer": answer, "document_id": did, "question": request.question, "ai_used": HUGGING_FACE_ENABLED}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Ошибка чата: {e}")
        return {"answer": "Произошла ошибка при обработке вопроса.", "document_id": request.document_id, "error": str(e)[:200]}

# ========== АУТЕНТИФИКАЦИЯ ==========
@app.post("/api/auth/login")
async def login_user(request: LoginRequest):
    try:
        logger.info(f"🔐 Вход пользователя: {request.email}")
        if not request.email or "@" not in request.email:
            raise HTTPException(status_code=400, detail="Invalid email format")
        if not request.password or len(request.password) < 1:
            raise HTTPException(status_code=400, detail="Password required")
        now = datetime.now().isoformat()
        return {
            "id": 1,
            "email": request.email,
            "username": request.email.split("@")[0],
            "token": f"demo_token_{int(datetime.now().timestamp())}",
            "created_at": now,
            "status": "success",
            "message": "Login successful",
            "timestamp": now,
            "demo_mode": True,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Ошибка логина: {e}")
        raise HTTPException(status_code=500, detail=f"Login error: {str(e)}")

@app.post("/api/auth/register")
async def register_user(request: RegisterRequest):
    try:
        logger.info(f"📝 Регистрация: {request.email} ({request.username})")
        if not request.email or "@" not in request.email:
            raise HTTPException(status_code=400, detail="Invalid email format")
        if not request.username or len(request.username) < 2:
            raise HTTPException(status_code=400, detail="Username must be at least 2 characters")
        if not request.password or len(request.password) < 1:
            raise HTTPException(status_code=400, detail="Password required")
        now = datetime.now().isoformat()
        return {
            "id": 2,
            "email": request.email,
            "username": request.username,
            "token": f"demo_token_{int(datetime.now().timestamp())}",
            "created_at": now,
            "status": "success",
            "message": "Registration successful",
            "timestamp": now,
            "demo_mode": True,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Ошибка регистрации: {e}")
        raise HTTPException(status_code=500, detail=f"Registration error: {str(e)}")

@app.get("/api/auth/check")
async def check_auth(token: str):
    try:
        if token and token.startswith("demo_token_"):
            return {"status": "success", "valid": True, "demo_mode": True, "timestamp": datetime.now().isoformat()}
        else:
            return {"status": "error", "valid": False, "message": "Invalid token format", "timestamp": datetime.now().isoformat()}
    except Exception as e:
        logger.error(f"❌ Ошибка проверки токена: {e}")
        raise HTTPException(status_code=500, detail=f"Auth check failed: {str(e)}")

@app.get("/api/auth/me")
async def get_current_user(authorization: str = ""):
    try:
        token = authorization.replace("Bearer ", "").strip() if authorization else ""
        if not token or not token.startswith("demo_token_"):
            raise HTTPException(status_code=401, detail="Токен недействителен")
        return {"id": 1, "email": "user@versevo.app", "username": "user", "token": token, "created_at": datetime.now().isoformat()}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Ошибка получения пользователя: {e}")
        raise HTTPException(status_code=500, detail=f"Auth error: {str(e)}")

@app.post("/api/auth/logout")
async def logout_user():
    return {"status": "success", "message": "Logout successful"}

# ========== ОТЧЁТЫ (mock) ==========
@app.get("/api/reports/mock/system-health")
async def report_system_health():
    return {"report_type": "system_health", "timestamp": datetime.now().isoformat(),
            "summary": {"total_users": 125, "active_users_7d": 58, "active_users_30d": 92, "retention_rate": 73.6},
            "table_statistics": {"users": 125, "documents": 89, "document_notes": 234, "document_analysis": 67, "favorite_quotes": 45},
            "is_mock": True}

@app.get("/api/reports/mock/user-activity")
async def report_user_activity():
    return {"report_type": "user_activity", "timestamp": datetime.now().isoformat(),
            "summary": {"total_users": 3, "active_users": 2, "total_documents": 25, "total_words_read": 258000, "activity_rate": 66.7},
            "data": [
                {"id": 1, "email": "admin@versevo.ru", "username": "admin", "documents_count": 12, "activity_status": "active"},
                {"id": 2, "email": "user@example.com", "username": "user1", "documents_count": 8, "activity_status": "active"},
                {"id": 3, "email": "inactive@test.com", "username": "old_user", "documents_count": 5, "activity_status": "inactive"}],
            "is_mock": True}

@app.get("/api/reports/mock/document-statistics")
async def report_document_statistics():
    return {"report_type": "document_statistics", "timestamp": datetime.now().isoformat(),
            "summary": {"total_documents": 3, "total_words": 686500, "avg_words": 228833.3, "languages_count": 2},
            "data": [
                {"filename": "Pride_and_Prejudice.pdf", "language": "en", "word_count": 125000, "reading_time_minutes": 625},
                {"filename": "Voyna_i_mir.txt", "language": "ru", "word_count": 560000, "reading_time_minutes": 2800},
                {"filename": "Les_Miserables.pdf", "language": "fr", "word_count": 530000, "reading_time_minutes": 2650}],
            "is_mock": True}

@app.get("/api/reports/mock/translation-usage")
async def report_translation_usage():
    return {"report_type": "translation_usage", "timestamp": datetime.now().isoformat(),
            "summary": {"total_translations": 23, "total_characters": 11700, "unique_translations": 19},
            "daily_data": [
                {"date": "2024-01-20", "translation_count": 15, "translation_service": "gemini"},
                {"date": "2024-01-21", "translation_count": 8, "translation_service": "gemini"},
                {"date": "2024-01-22", "translation_count": 0, "translation_service": "none"},
                {"date": "2024-01-23", "translation_count": 12, "translation_service": "openai"}],
            "is_mock": True}

@app.get("/api/reports/mock/ai-analysis")
async def report_ai_analysis():
    return {"report_type": "ai_analysis", "timestamp": datetime.now().isoformat(),
            "summary": {"total_analysis": 8, "unique_documents": 5},
            "sentiment_distribution": [
                {"sentiment": "Положительный", "percentage": 60.0},
                {"sentiment": "Нейтральный", "percentage": 25.0},
                {"sentiment": "Отрицательный", "percentage": 15.0}],
            "is_mock": True}

@app.get("/api/quotes/favorites")
async def get_favorite_quotes():
    return []

@app.post("/api/quotes/favorites")
async def add_favorite_quote(request: dict):
    return {"status": "success", "message": "Цитата добавлена в избранное"}

# ========== ЦИТАТЫ ==========
@app.get("/api/documents/{document_id}/quotes")
async def get_document_quotes(document_id: int, limit: int = 5):
    if document_id not in documents_store:
        raise HTTPException(status_code=404, detail="Document not found")
    content = documents_store[document_id].get("content", "")
    if not content or len(content.strip()) < 10:
        return {"document_id": document_id, "quotes": ["Текст документа пустой"], "count": 1, "fallback": True}
    sentences = re.split(r'(?<=[.!?])\s+', content)
    quotes = [s.strip() for s in sentences if 30 < len(s.strip()) < 250][:limit * 2]
    unique, seen = [], set()
    for q in quotes:
        norm = ' '.join(q.lower().split())
        if not any(len(set(norm.split()) & set(s.split())) / max(len(set(norm.split()) | set(s.split())), 1) > 0.7 for s in seen):
            unique.append(q)
            seen.add(norm)
        if len(unique) >= limit:
            break
    if len(unique) < limit:
        fallback = ["Каждая книга открывает новые горизонты.",
                    "Чтение — это диалог с автором через время и пространство.",
                    "Слова имеют силу менять восприятие мира.",
                    "Литература хранит мудрость поколений.",
                    "Текст — это мост между мыслью и её воплощением."]
        unique.extend(fallback[:limit - len(unique)])
    return {"document_id": document_id, "quotes": unique[:limit], "count": len(unique[:limit]), "fallback": len(unique[:limit]) > limit}

# ========== ЗАПУСК ==========
if __name__ == "__main__":
    import uvicorn
    logger.info(f"🚀 ЗАПУСК VERSION 5.2 НА ПОРТУ {PORT}")
    logger.info(f"🔤 Перевод: Hugging Face + Fallback")
    logger.info(f"🤖 HF перевод: {'ДОСТУПЕН' if hf_translator.is_available('en', 'ru') else 'НЕ ДОСТУПЕН'}")
    logger.info(f"📊 HF анализ: {'ДОСТУПЕН' if HUGGING_FACE_ENABLED else 'НЕ ДОСТУПЕН'}")
    logger.info(f"📈 NLTK анализ: {'ДОСТУПЕН' if NLTK_AVAILABLE else 'НЕ ДОСТУПЕН'}")
    uvicorn.run(app, host="0.0.0.0", port=PORT, log_level="info", access_log=True)
