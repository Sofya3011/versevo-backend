from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from typing import List, Optional
import base64
import uuid
import os
from datetime import datetime

from .db import get_db
from .models import Document
from .utils import detect_language_safe
from .config import settings

router = APIRouter(prefix="/api/flutter", tags=["flutter"])

def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Извлечение текста из различных форматов файлов"""
    try:
        import fitz  # PyMuPDF
        import docx
        
        if file_type == 'pdf':
            text = []
            try:
                doc = fitz.open(file_path)
                for page in doc:
                    text.append(page.get_text())
                doc.close()
            except Exception as e:
                return f"Ошибка чтения PDF: {str(e)}"
            return "\n\n".join(text)
        elif file_type in ['docx', 'doc']:
            try:
                doc = docx.Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n\n".join(paragraphs)
            except Exception as e:
                return f"Ошибка чтения DOCX: {str(e)}"
        elif file_type == 'txt':
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        return f"Ошибка извлечения текста: {str(e)}"

def detect_chapters(text: str) -> List[Dict]:
    """Автоматическое определение глав в тексте"""
    import re
    chapters = []
    
    patterns = [
        r'^(Глава\s+\d+[.:]\s*.+)$',
        r'^(CHAPTER\s+\d+[.:]\s*.+)$',
        r'^(Часть\s+\d+[.:]\s*.+)$',
        r'^(Part\s+\d+[.:]\s*.+)$',
        r'^(\d+[.:]\s*.+)$',
        r'^([IVXLCDM]+[.:]\s*.+)$'
    ]
    
    lines = text.split('\n')
    current_chapter = None
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        is_chapter = False
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                is_chapter = True
                break
                
        if is_chapter:
            if current_chapter:
                current_chapter['content'] = current_chapter['content'].strip()
                chapters.append(current_chapter)
            
            current_chapter = {
                'title': line,
                'start_position': sum(len(lines[j]) + 1 for j in range(i)),
                'content': ''
            }
        elif current_chapter:
            current_chapter['content'] += line + '\n'
    
    if current_chapter:
        current_chapter['content'] = current_chapter['content'].strip()
        chapters.append(current_chapter)
    
    if not chapters:
        chapters.append({
            'title': 'Основной текст',
            'start_position': 0,
            'content': text
        })
    
    return chapters

@router.get("/health")
async def health_check():
    """Эндпоинт для проверки здоровья приложения"""
    return {"status": "healthy", "service": "versevo-backend"}

@router.post("/upload")
async def flutter_upload(
    file: UploadFile = File(...),
    user_id: Optional[int] = Form(None),
    db: Session = Depends(get_db)
):
    """Загрузка файла из Flutter приложения"""
    try:
        # Сохраняем файл
        file_id = str(uuid.uuid4())
        file_extension = file.filename.split('.')[-1].lower() if '.' in file.filename else 'txt'
        file_path = f"{settings.UPLOAD_FOLDER}/{file_id}.{file_extension}"
        
        os.makedirs(settings.UPLOAD_FOLDER, exist_ok=True)
        
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        # Извлекаем текст
        content_str = extract_text_from_file(file_path, file_extension)
        
        # Определяем язык
        language = detect_language_safe(content_str)
        
        # Определяем главы
        chapters = detect_chapters(content_str)
        
        # Создаем документ в БД
        db_document = Document(
            filename=file.filename,
            content=content_str,
            language=language,
            file_type=file_extension,
            file_size=len(content),
            file_path=file_path,
            user_id=user_id if user_id else 1,  # TODO: заменить на реальный ID пользователя
            word_count=len(content_str.split()),
            char_count=len(content_str),
            chapter_count=len(chapters),
            reading_time_minutes=max(1, len(content_str.split()) // 200),
            metadata={
                "chapters": chapters,
                "original_filename": file.filename
            }
        )
        
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        
        return {
            "success": True,
            "document_id": db_document.id,
            "filename": db_document.filename,
            "language": db_document.language,
            "chapter_count": len(chapters),
            "word_count": db_document.word_count,
            "reading_time": db_document.reading_time_minutes
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.get("/documents")
async def get_user_documents(
    user_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    """Получение документов пользователя"""
    query = db.query(Document)
    
    if user_id:
        query = query.filter(Document.user_id == user_id)
    
    documents = query.order_by(Document.created_at.desc()).all()
    
    return [
        {
            "id": doc.id,
            "filename": doc.filename,
            "language": doc.language,
            "file_size": doc.file_size,
            "word_count": doc.word_count,
            "chapter_count": doc.chapter_count,
            "reading_time": doc.reading_time_minutes,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
        }
        for doc in documents
    ]

@router.get("/documents/{document_id}")
async def get_document_details(
    document_id: int,
    db: Session = Depends(get_db)
):
    """Получение деталей документа"""
    document = db.query(Document).filter(Document.id == document_id).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "id": document.id,
        "filename": document.filename,
        "content": document.content,
        "language": document.language,
        "file_type": document.file_type,
        "file_size": document.file_size,
        "chapters": document.metadata.get("chapters", []) if document.metadata else [],
        "word_count": document.word_count,
        "char_count": document.char_count,
        "reading_time": document.reading_time_minutes,
        "created_at": document.created_at.isoformat() if document.created_at else None
    }

@router.post("/analyze/{document_id}")
async def analyze_document(
    document_id: int,
    db: Session = Depends(get_db)
):
    """Анализ документа"""
    document = db.query(Document).filter(Document.id == document_id).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Простой анализ (можно расширить)
    analysis_result = {
        "summary": f"Документ '{document.filename}' содержит {document.word_count} слов, {document.chapter_count} глав.",
        "language": document.language,
        "chapter_count": document.chapter_count,
        "reading_time": f"{document.reading_time_minutes} минут",
        "complexity": "Сложный" if document.word_count > 5000 else "Средний" if document.word_count > 1000 else "Простой",
        "file_size": f"{document.file_size / 1024:.1f} KB" if document.file_size < 1024*1024 else f"{document.file_size / (1024*1024):.1f} MB"
    }
    
    return analysis_result
